# -*- coding: utf-8 -*-
"""
—————————————————————————————————————————
      @Author: OwObEEEEEE               
       @Email: 18092606532@163.com      
        @Time: 2023/12/8 11:50          
        @File: checkfile_creator.py
     @Version: v0.0.1              
    @Software: PyCharm          
 @Description:                          
—————————————————————————————————————————
"""
import os
import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# =====================================================
base_dir = os.path.dirname(os.path.abspath('Word_generator.py'))
file_path = os.path.join(base_dir, 'template.docx')


# =====================================================
# 创建
def create_table_p1_head(client, date, engineer, engineer_tel, contact=None, contact_tel=None):
    """

    :param client:
    :param date:
    :param contact:
    :param contact_tel:
    :param engineer:
    :param engineer_tel:
    """
    doc = docx.Document()

    # 创建表格
    table_p1_head = doc.add_table(rows=4, cols=4)

    # 合并第一行单元格
    merge_cells_horizontally(table_p1_head, 0, 0, 3)

    # 设置表格样式
    table_p1_head.style = 'Table Grid'
    table_p1_head.allow_autofit = False

    # 设置单元格内容和样式
    set_cell_content(table_p1_head.cell(0, 0), '{} 巡检日报'.format(client), bold=True, font_size=12)

    set_cell_content(table_p1_head.cell(1, 0), '客户名称', bold=True, font_size=10.5)
    set_cell_content(table_p1_head.cell(1, 1), client, font_size=10.5)
    set_cell_content(table_p1_head.cell(1, 2), '巡检日期', bold=True, font_size=10.5)
    set_cell_content(table_p1_head.cell(1, 3), date, font_size=10.5)

    set_cell_content(table_p1_head.cell(2, 0), '用户联系人', bold=True, font_size=10.5)
    set_cell_content(table_p1_head.cell(2, 1), contact, font_size=10.5)
    set_cell_content(table_p1_head.cell(2, 2), '用户联系人电话', bold=True, font_size=10.5)
    set_cell_content(table_p1_head.cell(2, 3), contact_tel, font_size=10.5)

    set_cell_content(table_p1_head.cell(3, 0), '巡检工程师', bold=True, font_size=10.5)
    set_cell_content(table_p1_head.cell(3, 1), engineer, font_size=10.5)
    set_cell_content(table_p1_head.cell(3, 2), '巡检工程师电话', bold=True, font_size=10.5)
    set_cell_content(table_p1_head.cell(3, 3), engineer_tel, font_size=10.5)

    # 保存word文档
    doc.save(os.path.join(base_dir, '{}_{}_巡检报告_{}.docx'.format(date, client, engineer)))


# =====================================================
# 合并单元格
def merge_cells_horizontally(table, row_index, col_start, col_end):
    """

    :param table: 表格
    :param row_index: 行号
    :param col_start: 起始列
    :param col_end: 终止列
    """
    row = table.rows[row_index]
    cell_start = row.cells[col_start]
    cell_end = row.cells[col_end]
    cell_start.merge(cell_end)


# =====================================================
# 设置 设置单元格的内容和样式
def set_cell_content(cell, content, bold=False, font_size=None):
    """

    :param cell:
    :param content:
    :param bold:
    :param font_size:
    """
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 垂直对齐方式设置为居中对齐
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落的水平对齐方式设置为居中对齐
    run = cell.paragraphs[0].add_run(content)  # 要在单元格中添加的文本内容
    run.bold = bold  # 设置粗体

    if font_size is not None:
        font = run.font
        font.size = Pt(font_size)


# =====================================================
# 隐藏外部作用域中的变量
def generate_inspection_report():
    client_name = input("请输入客户名称：")
    check_date = input("请输入巡检日期：")
    client_contact = input("请输入用户联系人（选填）：")
    # if client_contact == '':
    #     client_contact = " "
    client_contact_phone = input("请输入用户联系人电话（选填）：")
    # if client_contact_phone == '':
    #     client_contact_phone = " "
    engineer = input("请输入巡检工程师：")
    engineer_phone = input("请输入巡检工程师电话：")

    create_table_p1_head(client_name, check_date, engineer, engineer_phone, client_contact, client_contact_phone, )


# =====================================================


generate_inspection_report()